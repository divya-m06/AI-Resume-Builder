from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io

def generate_pdf(resume_data):
    """
    Generate a PDF resume from resume data.
    Returns bytes of the PDF file.
    """

    buffer = io.BytesIO()

    # PDF document setup
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=50, rightMargin=50,
        topMargin=50, bottomMargin=40
    )

    styles = getSampleStyleSheet()

    # Main title style
    title_style = ParagraphStyle(
        "TitleStyle",
        fontSize=26,
        leading=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#003b88"),
        spaceAfter=15,
        bold=True,
    )

    # Section title
    section_title = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=15,
        textColor=colors.HexColor("#003b88"),
        spaceBefore=10,
        spaceAfter=4,
    )

    # Normal text
    normal = ParagraphStyle(
        "Normal",
        parent=styles["Normal"],
        fontSize=11,
        leading=15,
    )

    # Small centered line below title
    def separator():
        return Table(
            [[" "]],
            colWidths=[460],
            style=[("LINEBELOW", (0, 0), (-1, -1), 1, colors.HexColor("#003b88"))]
        )

    elements = []

    # Title
    elements.append(Paragraph(resume_data["full_name"], title_style))
    elements.append(separator())
    elements.append(Spacer(1, 12))

    # Contact
    contact = f"{resume_data.get('email', '')}  |  {resume_data.get('phone', '')}"
    if contact.strip():
        elements.append(Paragraph(contact, normal))
        elements.append(Spacer(1, 15))

    # Job Role
    if resume_data.get("job_role"):
        elements.append(Paragraph("Professional Title", section_title))
        elements.append(Paragraph(resume_data["job_role"], normal))
        elements.append(Spacer(1, 10))

    # Education
    if resume_data.get("education"):
        elements.append(Paragraph("Education", section_title))
        for line in resume_data["education"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # Experience
    if resume_data.get("experience"):
        elements.append(Paragraph("Experience", section_title))
        for line in resume_data["experience"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # Skills
    skills = resume_data.get("skills", [])
    if not isinstance(skills, list):
        skills = skills.split(",")

    if skills:
        elements.append(Paragraph("Skills", section_title))
        elements.append(Paragraph(", ".join(skills), normal))
        elements.append(Spacer(1, 10))

    # Projects
    if resume_data.get("projects"):
        elements.append(Paragraph("Projects", section_title))
        for line in resume_data["projects"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # Previous Company
    if resume_data.get("previous_company"):
        elements.append(Paragraph("Previous Company Details", section_title))
        for line in resume_data["previous_company"].splitlines():
            elements.append(Paragraph("• " + line, normal))
        elements.append(Spacer(1, 10))

    # Border on all pages
    def add_border(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#003b88"))
        canvas.setLineWidth(2)
        canvas.roundRect(30, 30, 550, 735, radius=12)
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_border, onLaterPages=add_border)

    buffer.seek(0)
    return buffer.getvalue()

def generate_docx(resume_data):
    """
    Generate a DOCX resume from resume data.
    Returns bytes of the DOCX file.
    """

    doc = Document()

    # Border
    section = doc.sections[0]
    border_xml = f'''
        <w:pgBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="12" w:color="003b88"/>
            <w:left w:val="single" w:sz="12" w:color="003b88"/>
            <w:right w:val="single" w:sz="12" w:color="003b88"/>
            <w:bottom w:val="single" w:sz="12" w:color="003b88"/>
        </w:pgBorders>
    '''
    section._sectPr.append(parse_xml(border_xml))

    # Title
    title = doc.add_paragraph()
    run = title.add_run(resume_data["full_name"])
    run.bold = True
    run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact
    contact = f"{resume_data.get('email', '')} | {resume_data.get('phone', '')}"
    p = doc.add_paragraph(contact)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    # Helper to insert sections
    def add_section(title, content, bullet=False):
        doc.add_heading(title, level=2)
        for line in content:
            if bullet:
                doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(line)

    # Job Role
    if resume_data.get("job_role"):
        doc.add_heading("Professional Title", level=2)
        doc.add_paragraph(resume_data["job_role"])

    # Education
    if resume_data.get("education"):
        add_section("Education", resume_data["education"].splitlines(), True)

    # Experience
    if resume_data.get("experience"):
        add_section("Experience", resume_data["experience"].splitlines(), True)

    # Skills
    skills = resume_data.get("skills", [])
    if not isinstance(skills, list):
        skills = skills.split(",")
    if skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(skills))

    # Projects
    if resume_data.get("projects"):
        add_section("Projects", resume_data["projects"].splitlines(), True)

    # Previous Company
    if resume_data.get("previous_company"):
        add_section(
            "Previous Company Details",
            resume_data["previous_company"].splitlines()
        )

    # Export file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()